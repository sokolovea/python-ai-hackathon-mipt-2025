import re
import logging
import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

import subprocess
import platform
import shutil
from pathlib import Path
import logging

try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

# Настройка логгирования как в исходном коде
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)


class MarkdownConverter:
    def __init__(self, config_path: str = "config.json"):
        self.config = self.load_config(config_path)
        self.doc = Document()
        self.elements = []

    @staticmethod
    def load_config(path: str) -> dict:
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            logging.error(f"Файл конфигурации {path} не найден.")
            raise
        except json.JSONDecodeError as e:
            logging.error(f"Ошибка парсинга JSON в конфиге: {e}")
            raise
        except Exception as e:
            logging.error(f"Ошибка при загрузке конфига: {e}")
            raise

    def clean_text(self, text: str, remove_empty_lines: bool = True) -> str:
        lines = text.splitlines()
        cleaned_lines = [re.sub(r'\s+', ' ', line).strip() for line in lines]
        if remove_empty_lines:
            cleaned_lines = [line for line in cleaned_lines if line]
        return '\n'.join(cleaned_lines)

    def parse_markdown(self, markdown_text: str):
        elements = []
        lines = markdown_text.splitlines()
        i = 0
        prev_empty = False

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                if not prev_empty:
                    elements.append(("newline", ""))
                    prev_empty = True
                i += 1
                continue

            prev_empty = False

            if line.startswith("# "):
                elements.append(("heading1", self.clean_text(line[2:])))
            elif line.startswith("## "):
                elements.append(("heading2", self.clean_text(line[3:])))
            elif line.startswith("### "):
                elements.append(("heading3", self.clean_text(line[4:])))
            elif line.startswith("#### "):
                elements.append(("heading4", self.clean_text(line[5:]))) 
            elif line.startswith("##### "):
                elements.append(("heading5", self.clean_text(line[6:])))
            elif line.startswith("###### "):
                elements.append(("heading6", self.clean_text(line[7:])))   
            elif line.startswith("\\["):
                math_content = []
                while i < len(lines):
                    l = lines[i].strip()
                    if l.endswith("\\]"):
                        math_content.append(self.clean_text(l[:-2]))
                        break
                    math_content.append(self.clean_text(l[2:] if i == 0 else l))
                    i += 1
                elements.append(("block_math", "\n".join(math_content)))
            else:
                cleaned_line = self.clean_text(line)
                parts = re.split(r'(\\\(.*?\\\)|\$.*?\$)', cleaned_line)
                formatted_parts = []

                for part in parts:
                    if not part:
                        continue
                    if part.startswith("\\(") and part.endswith("\\)"):
                        formatted_parts.append(("inline_math", self.clean_text(part[2:-2])))
                    elif part.startswith("$") and part.endswith("$"):
                        formatted_parts.append(("inline_math", self.clean_text(part[1:-1])))
                    else:
                        text_parts = []
                        bold_parts = re.split(r'(\*\*.*?\*\*)', part)
                        for bp in bold_parts:
                            if bp.startswith("**") and bp.endswith("**"):
                                text_parts.append(("bold", self.clean_text(bp[2:-2])))
                            else:
                                italic_parts = re.split(r'(\*.*?\*)', bp)
                                for ip in italic_parts:
                                    if ip.startswith("*") and ip.endswith("*"):
                                        text_parts.append(("italic", self.clean_text(ip[1:-1])))
                                    else:
                                        text_parts.append(("plain", self.clean_text(ip)))
                        formatted_parts.extend(text_parts)

                elements.append(("paragraph", formatted_parts))
            i += 1

        self.elements = elements

    def set_styles(self):
        normal = self.doc.styles['Normal']
        normal.font.name = self.config["font"]
        normal.font.size = Pt(self.config["font_size"])

        para_format = normal.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para_format.line_spacing = Pt(1.5 * self.config["font_size"])
        para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para_format.first_line_indent = Inches(self.config["first_line_indent"])
        para_format.space_after = Pt(0)

        for level, size in self.config["heading_sizes"].items():
            style = f"Heading {level}"
            if style not in self.doc.styles:
                continue
            heading = self.doc.styles[style]
            heading.font.name = self.config["font"]
            heading.font.size = Pt(size)
            heading.font.bold = True
            heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            heading.paragraph_format.first_line_indent = Inches(0)

    def add_math_paragraph(self, math_text, is_inline=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if not is_inline else WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_before = Pt(6 if not is_inline else 0)
        p.paragraph_format.space_after = Pt(6 if not is_inline else 0)

        run = p.add_run(math_text)
        run.font.name = "Cambria Math"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
        run.font.size = Pt(self.config["font_size"])

    def build_doc(self):
        self.set_styles()
        for section in self.doc.sections:
            section.top_margin = Inches(self.config["margins"]["top"])
            section.bottom_margin = Inches(self.config["margins"]["bottom"])
            section.left_margin = Inches(self.config["margins"]["left"])
            section.right_margin = Inches(self.config["margins"]["right"])

        for elem_type, content in self.elements:
            if elem_type.startswith("heading"):
                level = int(elem_type[-1])
                self.doc.add_heading(content, level=level)
            elif elem_type == "block_math":
                self.add_math_paragraph(content)
            elif elem_type == "paragraph":
                p = self.doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(self.config["first_line_indent"])
                for part_type, text in content:
                    run = p.add_run(text)
                    run.font.name = self.config["font"]
                    run.font.size = Pt(self.config["font_size"])
                    if part_type == "bold":
                        run.bold = True
                    elif part_type == "italic":
                        run.italic = True
                    elif part_type == "inline_math":
                        run.font.name = "Cambria Math"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
            elif elem_type == "newline":
                self.doc.add_paragraph()

    def save(self, output_path: str):
        try:
            self.doc.save(output_path)
            logging.info(f"Документ сохранён: {output_path}")
        except Exception as e:
            logging.error(f"Ошибка при сохранении DOCX: {e}")

    def export_pdf(self, docx_path: str, pdf_path: str = None):
        system = platform.system()
        docx_path = Path(docx_path)
        pdf_path = Path(pdf_path) if pdf_path else docx_path.with_suffix(".pdf")

        if system == "Windows":
            try:
                from docx2pdf import convert as docx2pdf_convert
                docx2pdf_convert(str(docx_path), str(pdf_path))
                logging.info(f"PDF экспортирован (Windows/docx2pdf): {pdf_path}")
            except ImportError:
                logging.warning("Модуль docx2pdf не установлен — установите его для PDF экспорта на Windows")
            except Exception as e:
                logging.error(f"Ошибка при экспорте PDF через docx2pdf: {e}")
        elif system in {"Linux", "Darwin"}:
            libreoffice_path = shutil.which("libreoffice")
            if not libreoffice_path:
                logging.warning("LibreOffice не установлен или не найден в PATH — экспорт в PDF невозможен")
                return
            try:
                subprocess.run([
                    libreoffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    str(docx_path),
                    "--outdir", str(pdf_path.parent)
                ], check=True)
                logging.info(f"PDF экспортирован (LibreOffice): {pdf_path}")
            except subprocess.CalledProcessError as e:
                logging.error(f"Ошибка при экспорте PDF через LibreOffice: {e}")
        else:
            logging.warning(f"Платформа {system} не поддерживается для экспорта PDF")

    def convert(self, input_md: str, output_docx: str, export_pdf: bool = False):
        """
        Конвертация Markdown -> DOCX (и PDF при need).
        Если в config.use_pandoc == True, вызывает pandoc, иначе
        использует старую логику python-docx + LibreOffice.
        """
        logging.info(f"Чтение Markdown из {input_md}")
        try:
            md_text = Path(input_md).read_text(encoding="utf-8")
        except FileNotFoundError:
            logging.error(f"Файл {input_md} не найден.")
            return
        except Exception as e:
            logging.error(f"Ошибка при чтении Markdown: {e}")
            return

        if self.config.get("use_pandoc", False):
            tmp_md = Path(output_docx).with_suffix(".tmp.md")
            try:
                tmp_md.write_text(md_text, encoding="utf-8")

                cmd = [
                    "pandoc", str(tmp_md),
                    "-o", output_docx,
                    "--mathml"
                ]
                logging.info(f"Запуск Pandoc: {' '.join(cmd)}")
                subprocess.run(cmd, check=True)

                logging.info(f"DOCX успешно сгенерирован через Pandoc: {output_docx}")

                if export_pdf:
                    pdf_path = Path(output_docx).with_suffix(".pdf")
                    cmd_pdf = [
                        "pandoc", str(tmp_md),
                        "-o", str(pdf_path),
                        "--pdf-engine=xelatex",
                        "-V", "lang=ru",                   
                        "-V", "mainfont=DejaVu Serif",     
                        "-V", "sansfont=DejaVu Sans",      
                        "-V", "monofont=DejaVu Sans Mono"  
                    ]
                    logging.info(f"Запуск Pandoc для PDF: {' '.join(cmd_pdf)}")
                    subprocess.run(cmd_pdf, check=True)
                    logging.info(f"PDF успешно сгенерирован через Pandoc: {pdf_path}")

            except subprocess.CalledProcessError as e:
                logging.error(f"Pandoc вернул ненулевой код: {e.returncode}")
            except Exception as e:
                logging.error(f"Ошибка во время работы Pandoc: {e}")
            finally:
                try:
                    tmp_md.unlink()
                except Exception:
                    pass
        else:
            self.parse_markdown(md_text)
            self.build_doc()
            self.save(output_docx)
            if export_pdf:
                self.export_pdf(output_docx)

        logging.info("Конвертация завершена.") 

