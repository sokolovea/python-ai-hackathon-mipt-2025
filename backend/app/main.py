# backend/main.py
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware

import uuid
import subprocess

import numpy as np
from concurrent.futures import ThreadPoolExecutor, as_completed

from gigachat import GigaChat
from gigachat.models import Chat, Messages, MessagesRole

import logging
from tqdm import tqdm 
import time
import textwrap

from typing import List, Dict
from pathlib import Path

import os
import json

import hashlib

import re
import speech_recognition as sr
from moviepy.video.io.VideoFileClip import VideoFileClip
from pydub import AudioSegment
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


API_KEY = "NmE2MWQ5MTctOTQ3ZC00ZGI5LWIwODMtMGIxOTNkY2FiYzI5OjI3NTJjMzQwLWQ1NTMtNGJlNy1iMzI4LWE3YTMyYTA2NTNmYQ=="  
MAX_RETRIES = 3
DELAY_BETWEEN_REQUESTS = 90

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


logger = logging.getLogger(__name__)

giga = GigaChat(credentials=API_KEY, verify_ssl_certs=False)

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")
Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
Path(UPLOAD_DIR).mkdir(exist_ok=True)

def process_video(video_path: str, output_dir: str, interval: int = 10):
    audio_path = os.path.join(output_dir, "audio.wav")
    subprocess.run([
        "ffmpeg", "-i", video_path,
        "-q:a", "0", "-map", "a", audio_path
    ], check=True)

    frames_dir = os.path.join(output_dir, "frames")
    Path(frames_dir).mkdir(exist_ok=True)
    subprocess.run([
        "ffmpeg", "-i", video_path,
        "-vf", f"fps=1/{interval}",
        os.path.join(frames_dir, "frame_%04d.jpg")
    ], check=True)

@app.post("/upload/")
async def upload_video(file: UploadFile = File(...)):
    logging.info(f"Получен файл: {file.filename}")
    file_id = str(uuid.uuid4())
    output_dir = os.path.join(UPLOAD_DIR, file_id)
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    video_path = os.path.join(output_dir, "video.mp4")
    transcrib_path = os.path.join(output_dir, "transcrib.txt")
    TEXT_MD_PATH = os.path.join(output_dir, "summary.md")
    WORD_MD_PATH = os.path.join(output_dir, "summary.docx")
    try:
        with open(video_path, "wb") as buffer:
            buffer.write(await file.read())
        print(f"Видео сохранено: {video_path}")

        # process_video(video_path, output_dir)
        # text = start_transcribation(video_path)

        # with open(transcrib_path, 'w', encoding='UTF-8') as f:
        #     if isinstance(text, list):  
        #         f.writelines(text)
        #     else:  
        #         f.write(text)


        # source = LectureGenerator.load_source_text(transcrib_path)

        # generator = LectureGenerator(api_key=API_KEY, output_file=TEXT_MD_PATH)
        # generator.generate_lecture(source)
        # generator.save_to_file(TEXT_MD_PATH)


        # markdown_to_docx(transcrib_path, WORD_MD_PATH)

        text_md = None
        # with open(TEXT_MD_PATH, 'r', encoding='UTF-8') as f:
        #     text_md = f.readlines()
        
        return {"file_id": file_id, "text_md": "temp"}
    except Exception as e:
        print(f"Ошибка обработки видео: {e}")
        raise HTTPException(500, str(e))


@app.get("/frames/{file_id}")
def get_frames_info(file_id: str):
    frames_dir = os.path.join(UPLOAD_DIR, file_id, "frames")
    if not os.path.exists(frames_dir):
        raise HTTPException(404)
    
    frames = sorted(os.listdir(frames_dir))
    return {"frames": [f"frames/{f}" for f in frames]}

@app.get("/files/{file_id}/{path:path}")
def get_file(file_id: str, path: str):
    file_path = os.path.join(UPLOAD_DIR, file_id, path)
    logging.info(f"Запрашиваемый путь: {file_path}")  
    if not os.path.exists(file_path):
        logging.error(f"Файл не найден: {file_path}")
        raise HTTPException(404)
    return FileResponse(file_path)



logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

OUTPUT_AUDIO_PATH = "output_audio.wav"

def extract_audio_from_video(video_path, output_audio_path):
    """
    Извлекает аудиодорожку из видеофайла и сохраняет её в формате WAV.

    Args:
        video_path (str): Путь к исходному видеофайлу.
        output_audio_path (str): Путь для сохранения извлеченного аудиофайла.

    Returns:
        None
    """
    try:
        logging.info(f"Начало извлечения аудио из видео: {video_path}")
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(
            output_audio_path,
            codec="pcm_s16le",
            fps=16000,
            nbytes=2,
            ffmpeg_params=["-ac", "1"]
        )
        logging.info(f"Аудио успешно извлечено и сохранено в: {output_audio_path}")
    except Exception as e:
        logging.error(f"Ошибка при извлечении аудио из видео: {e}")


def split_audio(audio_path, chunk_length_ms=30000):
    """
    Разделяет аудиофайл на фрагменты заданной длины.

    Args:
        audio_path (str): Путь к исходному аудиофайлу.
        chunk_length_ms (int): Длина каждого фрагмента в миллисекундах (по умолчанию 30000 мс).

    Returns:
        list: Список путей к временным файлам с фрагментами аудио.
    """
    try:
        logging.info(f"Начало разделения аудио на фрагменты: {audio_path}")
        audio = AudioSegment.from_wav(audio_path)
        chunks = []
        for i in range(0, len(audio), chunk_length_ms):
            chunk = audio[i:i + chunk_length_ms]
            chunk_path = f"chunk_{i // chunk_length_ms}.wav"
            chunk.export(chunk_path, format="wav")
            chunks.append(chunk_path)
        logging.info(f"Аудио успешно разделено на {len(chunks)} фрагментов.")
        return chunks
    except Exception as e:
        logging.error(f"Ошибка при разделении аудио на фрагменты: {e}")
        return []


def recognize_chunk(chunk_path, language="ru-RU"):
    """
    Распознает речь в одном фрагменте аудио с использованием Google Speech Recognition API.

    Args:
        chunk_path (str): Путь к файлу с фрагментом аудио.
        language (str): Язык распознавания (по умолчанию "ru-RU").

    Returns:
        str: Распознанный текст. Если распознавание не удалось, возвращается пустая строка.
    """
    recognizer = sr.Recognizer()
    try:
        logging.info(f"Начало распознавания речи в фрагменте: {chunk_path}")
        with sr.AudioFile(chunk_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data, language=language)
            logging.info(f"Распознанный текст для фрагмента {chunk_path}: {text}")
            return text
    except sr.UnknownValueError:
        logging.warning(f"Речь не распознана в фрагменте: {chunk_path}")
        return ""
    except sr.RequestError as e:
        logging.error(f"Ошибка при обращении к API для фрагмента {chunk_path}: {e}")
        return ""


def transcribe_audio(audio_path):
    """
    Выполняет транскрибацию всего аудиофайла, разделяя его на фрагменты и распознавая каждый фрагмент параллельно.

    Args:
        audio_path (str): Путь к исходному аудиофайлу.

    Returns:
        tuple:
            - np.ndarray: Массив с распознанными текстами для каждого фрагмента.
            - str: Объединенный текст из всех фрагментов.
    """
    try:
        logging.info(f"Начало транскрибации аудио: {audio_path}")
        chunks = split_audio(audio_path)

        full_text = np.empty(len(chunks), dtype=object)

        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = {
                executor.submit(recognize_chunk, chunk, "ru-RU"): idx
                for idx, chunk in enumerate(chunks)
            }

            for future in as_completed(futures):
                idx = futures[future]
                try:
                    result = future.result()
                    full_text[idx] = result
                except Exception as e:
                    logging.error(f"Ошибка при обработке фрагмента {idx}: {e}")
                    full_text[idx] = ""

        for chunk in chunks:
            if os.path.exists(chunk):
                os.remove(chunk)
                logging.info(f"Временный файл удален: {chunk}")

        logging.info(f"Транскрибация завершена. Объединенный текст: {' '.join(full_text).strip()}")
        return full_text, " ".join(full_text).strip()
    except Exception as e:
        logging.error(f"Ошибка при выполнении транскрибации: {e}")
        return np.array([]), ""


def start_transcribation(video_path):
    """
    Основная функция программы, которая выполняет следующие шаги:
    1. Извлекает аудио из видеофайла.
    2. Выполняет транскрибацию извлеченного аудио.
    3. Выводит результаты и удаляет временные файлы.

    Returns:
        None
    """
    try:
        logging.info("Запуск программы.")

        extract_audio_from_video(video_path, OUTPUT_AUDIO_PATH)

        transcribe_list, transcription_results = transcribe_audio(OUTPUT_AUDIO_PATH)
        logging.info(f"Результаты транскрибации: {transcription_results}")

        if os.path.exists(OUTPUT_AUDIO_PATH):
            os.remove(OUTPUT_AUDIO_PATH)
            logging.info(f"Временный файл удален: {OUTPUT_AUDIO_PATH}")

        logging.info("Программа завершена.")

        return transcription_results
    except Exception as e:
        logging.critical(f"Критическая ошибка в программе: {e}")


'''
todo: разбить класс на подклассы!!!
'''
class LectureGenerator:
    def __init__(self, api_key: str, output_file):
        if not api_key:
            raise ValueError("API key обязателен")
        self.client = GigaChat(credentials=api_key, verify_ssl_certs=False)
        self.lecture_content: List[str] = []
        self.cache_dir = "lecture_cache"
        self.output_file = output_file
        os.makedirs(self.cache_dir, exist_ok=True)

    def _send_prompt(self, prompt: str) -> str:
        for attempt in range(MAX_RETRIES):
            try:
                response = self.client.chat(
                    Chat(
                        messages=[
                            Messages(
                                role=MessagesRole.USER,
                                content=prompt
                            )
                        ],
                        temperature=0.7,
                        max_tokens=8000
                    )
                )
                return response.choices[0].message.content
            except Exception as e:
                logger.warning(f"Ошибка при запросе (попытка {attempt + 1}): {e}")
                if attempt == MAX_RETRIES - 1:
                    logger.error("Превышено количество попыток. Прерывание.")
                    raise
                time.sleep(DELAY_BETWEEN_REQUESTS * (attempt + 1))

    def _get_cache_path(self, name: str, content: str | None = None) -> str:
        safe_name = name.replace(" ", "_")[:50]
        if content:
            hash_digest = hashlib.md5(content.encode('utf-8')).hexdigest()[:10]
            safe_name += f"_{hash_digest}"
        return os.path.join(self.cache_dir, f"{safe_name}.json")

    def _load_cache(self, name: str, context: str | None = None) -> str | None:
        path = self._get_cache_path(name, context)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                logger.info(f"Загружаю из кеша: {path}")
                return json.load(f)['content']
        return None

    def _save_cache(self, name: str, content: str, context: str | None = None):
        path = self._get_cache_path(name, context)
        with open(path, 'w', encoding='utf-8') as f:
            json.dump({'content': content}, f, ensure_ascii=False, indent=2)

    @staticmethod
    def load_source_text(file_path: str, retries: int = 3, delay: int = 5) -> str:
        for attempt in range(retries):
            try:
                logger.info(f"Пробую загрузить файл: {file_path} (попытка {attempt + 1})")
                f = open(file_path, encoding="utf-8")  
                text = f.read()
                f.close()
                if not text.strip():
                    raise ValueError("Файл пуст.")
                return text
            except Exception as e:
                logger.warning(f"Не удалось загрузить файл: {e}")
                if attempt == retries - 1:
                    logger.error("Файл не загружен после нескольких попыток.")
                    raise
                time.sleep(delay * (attempt + 1))


    def generate_lecture(self, source_text: str) -> None:
        logger.info("Запуск генерации структуры лекции...")

        if not source_text.strip():
            logger.error("Источник текста пуст. Прекращаю генерацию.")
            raise ValueError("Источник текста пуст. Пожалуйста, проверьте содержимое.")

        if len(source_text.strip()) < 100:
            logger.warning("Источник текста слишком короткий. Возможна низкая релевантность результата.")

        structure_cache = self._load_cache("structure", source_text)
        if structure_cache:
            structure = structure_cache
        else:
            structure_prompt = textwrap.dedent(f"""
            Ты — профессиональный преподаватель в ВУЗе. На основе следующего текста создай подробную структуру лекции в формате Markdown.
            1. **Не добавляй** фраз, не относящихся к теме, таких как: 
                - "Что-то в вашем вопросе меня смущает"
                - "Не люблю менять тему"
                - обращения к собеседнику ("давайте обсудим", "как вы думаете?" и т.п.)
            2. **Пиши от имени автора курса**, в официально-деловом или вдохновляющем стиле, без попыток вести диалог.
            3. Не вставляй шаблонные фразы, которые не связаны с содержанием (например, защитные реплики ИИ).
            4. Текст должен быть цельным, без повтора одних и тех же фраз в конце (проверяй дубли).
            Требования:
            - Начни с оглавления.
            - Включи логичные главы и подглавы (используй заголовки ## и ###).
            - Последняя глава — обязательно "Резюме" или "Выводы".
            - Названия разделов должны быть ёмкими, но не слишком короткими.
            - Не добавляй содержание самих разделов — только структуру.

            Текст лекции:
            {source_text}
            """)
            structure = self._send_prompt(structure_prompt)
            self._save_cache("structure", structure, source_text)

        self.lecture_content.append(structure)
        sections = self._extract_sections(structure)
        logger.info(f"Найдено {len(sections)} разделов. Приступаю к генерации содержания...")

        for section in tqdm(sections, desc="Генерация разделов"):
            title = section['title']
            content_cache = self._load_cache(title, source_text)
            if content_cache:
                self.lecture_content.append(content_cache)
                continue

            logger.info(f"Генерация раздела: {title}")

            content_prompt = textwrap.dedent(f"""
            Ты — экспертный преподаватель и технический писатель. Твоя задача — написать подробное и качественное содержание лекции по следующему разделу.

            Раздел: "{title}"

            Инструкции:
            - Используй академический стиль. Пиши от третьего лица. 
            - Применяй Markdown: заголовки, списки, таблицы, выделения, при необходимости.
            - Избегай "воды" и повторов, фокусируйся на сути.
            - Не выдумывай от себя — используй только информацию из контекста ниже.
            - Исправь ошибки или опечатки, если встретишь.
            - При необходимости можешь добавить примеры, схемы (опиши словами, например: *таблица с различиями X и Y*).

            Контекст:
            {source_text}
            """)

            content = self._send_prompt(content_prompt)
            self._save_cache(title, content,source_text)
            self.lecture_content.append(content)
            time.sleep(DELAY_BETWEEN_REQUESTS)

    def _extract_sections(self, markdown_text: str) -> List[Dict]:
        sections = []
        for line in markdown_text.splitlines():
            if line.startswith("## "):
                level = 2
            elif line.startswith("### "):
                level = 3
            else:
                continue
            sections.append({"title": line.strip("# ").strip(), "level": level})
        return sections

    def save_to_file(self, filename) -> None:
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(self.lecture_content))
            logger.info(f"Лекция успешно сохранена в файл: {filename}")
        except Exception as e:
            logger.error(f"Не удалось сохранить лекцию: {e}")

'''
---
'''


logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

def clean_text(text: str, remove_empty_lines: bool = True) -> str:
    """
    Очищает текст от лишних пробелов:
    - заменяет множественные пробелы на один
    - удаляет пробелы в начале/конце строк
    - удаляет пустые строки (если remove_empty_lines=True)
    """
    lines = text.splitlines()
    cleaned_lines = [re.sub(r'\s+', ' ', line).strip() for line in lines]
    if remove_empty_lines:
        cleaned_lines = [line for line in cleaned_lines if line]
    return '\n'.join(cleaned_lines)

def parse_markdown(markdown_text: str) -> list:
    """Парсер Markdown с очисткой пробелов"""
    elements = []
    lines = markdown_text.splitlines()
    i = 0
    prev_empty = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            if not prev_empty:
                elements.append(("newline", ""))
                prev_empty = True
            i += 1
            continue

        prev_empty = False

        if line.startswith("# "):
            elements.append(("heading1", clean_text(line[2:])))
        elif line.startswith("## "):
            elements.append(("heading2", clean_text(line[3:])))
        elif line.startswith("### "):
            elements.append(("heading3", clean_text(line[4:])))
        elif line.startswith("#### "):
            elements.append(("heading4", clean_text(line[5:])))
        elif line.startswith("##### "):
            elements.append(("heading5", clean_text(line[6:])))
        elif line.startswith("\\["):
            math_content = []
            while i < len(lines):
                l = lines[i].strip()
                if l.endswith("\\]"):
                    math_content.append(clean_text(l[:-2]))
                    break
                math_content.append(clean_text(l[2:] if i == 0 else l))
                i += 1
            elements.append(("block_math", "\n".join(math_content)))
        else:
            cleaned_line = clean_text(line)
            parts = re.split(r'(\\\(.*?\\\)|\$.*?\$)', cleaned_line)
            formatted_parts = []

            for part in parts:
                if not part:
                    continue
                if part.startswith("\\(") and part.endswith("\\)"):
                    formatted_parts.append(("inline_math", clean_text(part[2:-2])))
                elif part.startswith("$") and part.endswith("$"):
                    formatted_parts.append(("inline_math", clean_text(part[1:-1])))
                else:
                    text_parts = []
                    bold_parts = re.split(r'(\*\*.*?\*\*)', part)
                    for bp in bold_parts:
                        if bp.startswith("**") and bp.endswith("**"):
                            text_parts.append(("bold", clean_text(bp[2:-2])))
                        else:
                            italic_parts = re.split(r'(\*.*?\*)', bp)
                            for ip in italic_parts:
                                if ip.startswith("*") and ip.endswith("*"):
                                    text_parts.append(("italic", clean_text(ip[1:-1])))
                                else:
                                    text_parts.append(("plain", clean_text(ip)))
                    formatted_parts.extend(text_parts)

            elements.append(("paragraph", formatted_parts))
        i += 1

    logging.debug(f"Распознано {len(elements)} элементов")
    return elements

def set_doc_styles(doc):
    """
    Применяет стили оформления ко всему документу:
    
    - Основной стиль ('Normal'): шрифт Arial, размер 12 pt, выравнивание по ширине,
      полуторный межстрочный интервал, отступ первой строки.
    
    - Заголовки уровней 1–3: шрифт Arial, увеличенный размер шрифта, полужирный,
      выравнивание по левому краю, без отступа первой строки и без дополнительных отступов до/после.

    Используется для стандартизации внешнего вида документа и улучшения читаемости.
    """
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(12)

    para_format = normal_style.paragraph_format
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para_format.line_spacing = Pt(18)  # 1.5 * 12pt
    para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para_format.first_line_indent = Inches(0.25)
    para_format.space_after = Pt(0)

    heading_levels = {
        1: Pt(16),
        2: Pt(14),
        3: Pt(12)
    }

    for level, size in heading_levels.items():
        style_name = f'Heading {level}'
        if style_name not in doc.styles:
            continue  

        heading = doc.styles[style_name]
        heading_font = heading.font
        heading_font.name = 'Arial'
        heading_font.size = size
        heading_font.bold = True

        heading_format = heading.paragraph_format
        heading_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_format.first_line_indent = Inches(0)
        heading_format.space_before = Pt(0)
        heading_format.space_after = Pt(6)  


def add_math_paragraph(doc, math_text, is_inline=False):
    """Добавление формул"""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if not is_inline else WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_before = Pt(6 if not is_inline else 0)
    p.paragraph_format.space_after = Pt(6 if not is_inline else 0)

    run = p.add_run(math_text)
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(12)

def markdown_to_docx(md_path: str, docx_path: str):
    """Конвертация Markdown -> DOCX"""
    logging.info(f"Чтение markdown из {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    elements = parse_markdown(md_text)
    doc = Document()
    set_doc_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    for elem_type, content in elements:
        if elem_type.startswith("heading"):
            level = int(elem_type[-1])
            doc.add_heading(content, level=level)
        elif elem_type == "block_math":
            add_math_paragraph(doc, content)
        elif elem_type == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.25)
            for part_type, part_text in content:
                run = p.add_run(part_text)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                if part_type == "bold":
                    run.bold = True
                elif part_type == "italic":
                    run.italic = True
                elif part_type == "inline_math":
                    run.font.name = "Cambria Math"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
        elif elem_type == "newline":
            doc.add_paragraph()

    doc.save(docx_path)
    logging.info(f"Документ успешно сохранён в {docx_path}")
